#!/usr/bin/env python3
"""
Render a report written in Markdown to DOCX and PDF.

Markdown -> pandoc -> DOCX -> LibreOffice -> PDF.

The DOCX reference style is built at run time from pandoc's own default reference
document, so nothing binary lives in the repository and the styling never goes stale.
"""

import argparse
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

BODY_FONT = "Arial"
BODY_PT = 11
H1_COLOR = "1F3864"
H2_COLOR = "2E4D7B"
SUBTLE = "595959"

SOFFICE_CANDIDATES = [
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/usr/lib/libreoffice/program/soffice",
    "/opt/libreoffice/program/soffice",
]


def need(binary: str, hint: str) -> str:
    path = shutil.which(binary)
    if not path:
        sys.exit(f"{binary} not found. {hint}")
    return path


def find_soffice() -> str | None:
    path = shutil.which("soffice") or shutil.which("libreoffice")
    if path:
        return path
    for candidate in SOFFICE_CANDIDATES:
        if Path(candidate).exists():
            return candidate
    return None


def build_reference_doc(pandoc: str, dest: Path) -> Path | None:
    """Pandoc's default reference.docx, restyled. Returns None if python-docx is absent."""
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt, RGBColor
    except ImportError:
        return None

    raw = subprocess.run([pandoc, "--print-default-data-file", "reference.docx"],
                         capture_output=True, check=True).stdout
    dest.write_bytes(raw)
    document = docx.Document(str(dest))

    def restyle(name, *, size=None, color=None, bold=None):
        try:
            style = document.styles[name]
        except KeyError:
            return
        font = style.font
        font.name = BODY_FONT
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        # Theme fonts outrank explicit ones, so clear them before naming the family.
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
            if rfonts.get(qn(attr)) is not None:
                del rfonts.attrib[qn(attr)]
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), BODY_FONT)
        if size is not None:
            font.size = Pt(size)
        if color is not None:
            font.color.rgb = RGBColor.from_string(color)
        if bold is not None:
            font.bold = bold

    restyle("Normal", size=BODY_PT)
    restyle("Body Text", size=BODY_PT)
    restyle("Title", size=20, color=H1_COLOR, bold=True)
    restyle("Subtitle", size=12, color=H2_COLOR)
    restyle("Heading 1", size=14, color=H1_COLOR, bold=True)
    restyle("Heading 2", size=12, color=H2_COLOR, bold=True)
    restyle("Heading 3", size=11, color=H2_COLOR, bold=True)
    restyle("Author", size=10, color=SUBTLE)
    restyle("Date", size=10, color=SUBTLE)

    # Bordered table with a shaded header row, matching the report house style.
    try:
        table_style = document.styles["Table"]
    except KeyError:
        table_style = None
    if table_style is not None:
        tbl_pr = table_style.element.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table_style.element.append(tbl_pr)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            line = OxmlElement(f"w:{edge}")
            line.set(qn("w:val"), "single")
            line.set(qn("w:sz"), "4")
            line.set(qn("w:color"), "BFBFBF")
            borders.append(line)
        tbl_pr.append(borders)
        margins = OxmlElement("w:tblCellMar")
        for edge, width in (("top", 60), ("left", 100), ("bottom", 60), ("right", 100)):
            cell = OxmlElement(f"w:{edge}")
            cell.set(qn("w:w"), str(width))
            cell.set(qn("w:type"), "dxa")
            margins.append(cell)
        tbl_pr.append(margins)
        band = OxmlElement("w:tblStylePr")
        band.set(qn("w:type"), "firstRow")
        band_pr = OxmlElement("w:tcPr")
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "E8EEF4")
        band_pr.append(shd)
        band.append(band_pr)
        table_style.element.append(band)

    section = document.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    for side in ("left", "right", "top", "bottom"):
        setattr(section, f"{side}_margin", Mm(19))
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.text = ""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(SUBTLE)
    fld = run._r.makeelement(qn("w:fldSimple"), {qn("w:instr"): "PAGE"})
    run._r.addnext(fld)

    document.save(str(dest))
    return dest


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("markdown", help="report written in Markdown")
    ap.add_argument("--out", "-o", default=None,
                    help="output path stem (default: alongside the Markdown, same stem)")
    ap.add_argument("--no-pdf", action="store_true", help="stop after the DOCX")
    args = ap.parse_args()

    md = Path(args.markdown).expanduser().resolve()
    if not md.exists():
        sys.exit(f"not found: {md}")

    stem = Path(args.out).expanduser().resolve() if args.out else md.with_suffix("")
    stem.parent.mkdir(parents=True, exist_ok=True)
    docx_path = stem.with_suffix(".docx")

    pandoc = need("pandoc", "Install it: brew install pandoc")

    with tempfile.TemporaryDirectory(prefix="mar_render_") as tmp:
        reference = build_reference_doc(pandoc, Path(tmp) / "reference.docx")
        cmd = [pandoc, str(md),
               "--from", "markdown+pipe_tables+auto_identifiers",
               "--to", "docx", "--standalone", "-o", str(docx_path)]
        if reference:
            cmd += [f"--reference-doc={reference}"]
        else:
            print("note: python-docx not available — using pandoc default styling",
                  file=sys.stderr)
        subprocess.run(cmd, check=True)

    print(f"DOCX: {docx_path}")
    if args.no_pdf:
        return 0

    soffice = find_soffice()
    if not soffice:
        print("LibreOffice not found — DOCX written, PDF skipped.\n"
              "  Install it (brew install --cask libreoffice) and re-run,\n"
              "  or export the PDF from Word.", file=sys.stderr)
        return 1

    subprocess.run([soffice, "--headless", "--convert-to", "pdf",
                    "--outdir", str(docx_path.parent), str(docx_path)],
                   check=True, capture_output=True)
    pdf_path = docx_path.with_suffix(".pdf")
    if not pdf_path.exists():
        print("LibreOffice reported success but no PDF appeared.", file=sys.stderr)
        return 1
    print(f"PDF:  {pdf_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
